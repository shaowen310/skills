#!/usr/bin/env python3
"""
Convert Word (.docx) files to Markdown with image extraction.

Usage:
    python docx2md.py <input.docx> [output.md]

The output Markdown is written to *output.md* (defaults to <input_stem>.md).
Extracted images are placed in an ``assets/`` folder next to the output file.
EMF/WMF vector images are converted to PNG (requires ``pillow-emf``) so they
render in Markdown viewers.

For programmatic use, import the ``convert()`` function:
    from docx2md import convert
    md_path, by_page = convert("input.docx", "out.md", asset_dir="assets")
"""

from __future__ import annotations

import argparse
import ctypes
import importlib
import json
import re
import string
import sys
import unicodedata
from pathlib import Path
from typing import Any, final

try:
    from docx import Document
    from docx.document import Document as _Document
except ImportError:
    sys.exit("Missing dependency: python-docx. Run: pip install python-docx")

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None  # type: ignore[assignment]


# XML namespaces used for image detection (DrawingML + VML + relationships)
_REL_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
_DRAW_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_VML_NS = "{urn:schemas-microsoft-com:vml}"


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------

_CONTROL_RE = re.compile(r"[\u0000-\u0008\u000b\u000c\u000e-\u001f\u0080-\u009f]")
_FORMAT_CHARS_RE = re.compile(
    r"[\u200b-\u200f\u2028-\u202f\ufeff]"
)


def _clean_text(text: str) -> str:
    """Remove invisible / control characters and normalize whitespace.

    Word line breaks (``<w:br/>`` / ``<w:cr/>``) and soft wraps inside table
    cells are exposed by python-docx as literal ``\\n`` / ``\\r`` characters.
    Left in place they split a Markdown table row (or any block) across
    several physical lines, which breaks GFM table parsing. We convert them
    to ordinary spaces and then collapse runs of spaces.
    """
    if not text:
        return ""
    text = text.replace("\u000b", " ").replace("\u000c", " ").replace("\u00a0", " ")
    # Normalize newlines (from Word line breaks / cell soft wraps) to spaces
    # so they never break a Markdown table row onto multiple physical lines.
    text = text.replace("\r\n", " ").replace("\r", " ").replace("\n", " ")
    text = _CONTROL_RE.sub("", text)
    text = _FORMAT_CHARS_RE.sub("", text)
    text = re.sub(r"  +", " ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Inline formatting conversion
# ---------------------------------------------------------------------------

# ASCII punctuation characters, per CommonMark's definition
_ASCII_PUNCT = frozenset(string.punctuation)


def _is_punct_or_space(ch: str) -> bool:
    """True if *ch* is whitespace or a Unicode punctuation character.

    Used to decide whether bordering characters could break CommonMark's
    delimiter *flanking* rules. When the wrapped text begins/ends with such a
    character and the marker is not bordered by whitespace/punctuation in the
    surrounding text, ``**``/``*`` may be dropped. A single space inserted
    *outside* the markers (the "add a space after punctuation" scheme) keeps
    the emphasis valid
    while leaving the punctuation *inside* the emphasis for fidelity.
    """
    if ch.isspace():
        return True
    if ch in _ASCII_PUNCT:
        return True
    return unicodedata.category(ch).startswith("P")


def _wrap_inline(
    text: str,
    bold: bool,
    italic: bool,
    code_style: bool,
    href: str | None,
    keep_emphasis: bool = True,
) -> str:
    """Wrap *text* with the matching inline markers, keeping every character
    (including bordering punctuation) *inside* the emphasis for fidelity.

    CommonMark's delimiter *flanking* rules can drop ``**``/``*`` when the
    wrapped text begins/ends with a punctuation character and the marker is
    not bordered by whitespace/punctuation in the surrounding text. To keep
    the emphasis valid without moving punctuation outside (which would strip
    the original bold/italic styling from that punctuation), we add a single
    space *outside* the markers — but only when the border character is
    punctuation/space and     could therefore break flanking (the "add a space after punctuation"
    scheme).
    """
    if not text:
        return text

    if code_style:
        wrapped = f"`{text}`"
        emphasized = False
    elif bold and italic and keep_emphasis:
        wrapped = f"***{text}***"
        emphasized = True
    elif bold and keep_emphasis:
        wrapped = f"**{text}**"
        emphasized = True
    elif italic and keep_emphasis:
        wrapped = f"*{text}*"
        emphasized = True
    else:
        wrapped = text
        emphasized = False

    # "Add space after punctuation" scheme: only when the text starts/ends with
    # punctuation (or whitespace) that could break flanking do we add a single
    # space *outside* the markers. When a hyperlink is present, the surrounding
    # brackets [ ] are already punctuation and satisfy flanking, so no space is
    # added; code spans and plain text need no handling.
    if emphasized and not href:
        if _is_punct_or_space(text[0]):
            wrapped = " " + wrapped
        if _is_punct_or_space(text[-1]):
            wrapped = wrapped + " "

    if href:
        wrapped = f"[{wrapped}]({href})"
    return wrapped


def _run_style(run: Any) -> tuple[bool, bool, bool, str | None]:
    """Return the inline style key for *run* (used to merge like-formatted runs)."""
    bold = run.bold
    italic = run.italic
    # ``font.name`` may be None; treat as no special formatting
    code_style = run.font.name == "Courier New" if run.font.name else False
    href = _run_hyperlink_target(run)
    return (bold, italic, code_style, href)


def _runs_markdown(runs: list[Any], keep_emphasis: bool = True) -> str:
    """Join paragraph runs into Markdown, merging consecutive runs that share
    identical inline formatting so adjacent bold runs aren't double-wrapped."""
    segments: list[list[Any]] = []  # [[style_key, accumulated_text], ...]
    for run in runs:
        text = _clean_text(run.text)
        if not text:
            continue  # empty run: drop, does not break merge boundary
        key = _run_style(run)
        if segments and segments[-1][0] == key:
            segments[-1][1] += text
        else:
            segments.append([key, text])
    return "".join(
        _wrap_inline(t, *k, keep_emphasis=keep_emphasis) for k, t in segments
    )


def _run_hyperlink_target(run: Any) -> str | None:
    """Return the hyperlink target URL for *run*, or ``None``.

    python-docx does not expose a ``Run.hyperlink`` attribute, so we walk
    the run's XML ancestors to find a ``<w:hyperlink>`` element and resolve
    its relationship id against the document's relationships.
    """
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rel_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    element = run._element
    while element is not None:
        if element.tag == f"{ns}hyperlink":
            rid = element.get(f"{rel_ns}id")
            if rid:
                try:
                    rel = run._parent.part.rels[rid]
                    return rel.target_ref
                except Exception:  # noqa: BLE001
                    return None
        element = element.getparent()
    return None


# ---------------------------------------------------------------------------
# List processing
# ---------------------------------------------------------------------------

def _list_level(paragraph: Any) -> int:
    """Return the list nesting level (0-based) or -1 if not a list item."""
    numPr = paragraph._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
    if numPr is None:
        return -1
    ilvl = numPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
    if ilvl is not None and ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") is not None:
        return int(ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"))
    return 0


def _is_numbered_list(paragraph: Any) -> bool:
    """Heuristic: if ``numPr`` → ``numFmt`` is ``decimal`` or ``lowerLetter``,
    treat as numbered list, otherwise bullet."""
    numPr = paragraph._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
    if numPr is None:
        return False
    numFmt = numPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt")
    if numFmt is not None:
        val = numFmt.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "")
        return val in ("decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman")
    return False


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def _convert_vector_to_png(src: Path, dst: Path) -> bool:
    """Convert an EMF/WMF vector image to PNG. Returns ``True`` on success.

    Dispatch by format:
      * EMF  → Windows GDI (``ctypes``, no extra dependency) on win32;
               otherwise try ``pillow-emf`` if installed.
      * WMF  → Pillow's native WMF support, then ``PyMuPDF`` (``fitz``).

    Falls back gracefully (returns ``False``) when no converter is available so
    the caller can keep the original vector file.
    """
    suffix = src.suffix.lower()
    if suffix == ".emf":
        if sys.platform == "win32":
            if _emf_to_png_gdi(src, dst):
                return True
        if PILImage is not None:
            try:
                _ = importlib.import_module("pillow_emf")
                img = PILImage.open(src)
                img.convert("RGB").save(dst, "PNG")
                return True
            except Exception:  # noqa: BLE001
                pass
        return _wmf_or_emf_via_fitz(src, dst)
    # WMF
    if PILImage is not None:
        try:
            img = PILImage.open(src)
            img.convert("RGB").save(dst, "PNG")
            return True
        except Exception:  # noqa: BLE001
            pass
    return _wmf_or_emf_via_fitz(src, dst)


def _wmf_or_emf_via_fitz(src: Path, dst: Path) -> bool:
    """Last-resort EMF/WMF → PNG via PyMuPDF (limited format support)."""
    try:
        fitz = importlib.import_module("fitz")

        doc = fitz.open(str(src))
        if doc.page_count == 0:
            return False
        pix = doc[0].get_pixmap(matrix=fitz.Matrix(2, 2))
        pix.save(str(dst))
        doc.close()
        return True
    except Exception:  # noqa: BLE001
        return False


def _emf_to_png_gdi(src: Path, dst: Path, scale: int = 2) -> bool:
    """Render an EMF to PNG using the Windows GDI API (no 3rd-party deps).

    Uses ``GetEnhMetaFileW`` + ``PlayEnhMetaFile`` into a DIB section, then
    saves the bits via Pillow. Returns ``True`` on success.
    """
    if PILImage is None:
        return False
    try:
        gdi32 = ctypes.WinDLL("gdi32", use_last_error=True)  # type: ignore[name-defined]
        user32 = ctypes.WinDLL("user32", use_last_error=True)  # type: ignore[name-defined]
    except Exception:  # noqa: BLE001
        return False

    # Configure argument/return types so ctypes passes the opaque handles
    # (which appear as large negative Python ints) correctly.
    HENHMETAFILE = ctypes.c_void_p
    HDC = ctypes.c_void_p
    HGDIOBJ = ctypes.c_void_p
    HBRUSH = ctypes.c_void_p
    LPRECT = ctypes.c_void_p
    gdi32.GetEnhMetaFileW.restype = HENHMETAFILE
    gdi32.GetEnhMetaFileW.argtypes = [ctypes.c_wchar_p]
    gdi32.GetEnhMetaFileHeader.restype = ctypes.c_uint
    gdi32.GetEnhMetaFileHeader.argtypes = [HENHMETAFILE, ctypes.c_uint, ctypes.c_void_p]
    gdi32.CreateCompatibleDC.restype = HDC
    gdi32.CreateCompatibleDC.argtypes = [HDC]
    gdi32.CreateDIBSection.restype = HGDIOBJ
    gdi32.CreateDIBSection.argtypes = [
        HDC, ctypes.c_void_p, ctypes.c_uint, ctypes.c_void_p, ctypes.c_void_p, ctypes.c_uint
    ]
    gdi32.SelectObject.restype = HGDIOBJ
    gdi32.SelectObject.argtypes = [HDC, HGDIOBJ]
    gdi32.GetStockObject.restype = HGDIOBJ
    gdi32.GetStockObject.argtypes = [ctypes.c_int]
    gdi32.PlayEnhMetaFile.restype = ctypes.c_int
    gdi32.PlayEnhMetaFile.argtypes = [HDC, HENHMETAFILE, LPRECT]
    gdi32.GetDIBits.restype = ctypes.c_int
    gdi32.GetDIBits.argtypes = [
        HDC, HGDIOBJ, ctypes.c_uint, ctypes.c_uint, ctypes.c_void_p, ctypes.c_void_p, ctypes.c_uint
    ]
    gdi32.DeleteObject.restype = ctypes.c_int
    gdi32.DeleteObject.argtypes = [HGDIOBJ]
    gdi32.DeleteDC.restype = ctypes.c_int
    gdi32.DeleteDC.argtypes = [HDC]
    gdi32.DeleteEnhMetaFile.restype = ctypes.c_int
    gdi32.DeleteEnhMetaFile.argtypes = [HENHMETAFILE]
    user32.FillRect.restype = ctypes.c_int
    user32.FillRect.argtypes = [HDC, LPRECT, HBRUSH]

    @final
    class RECTL(ctypes.Structure):
        _fields_ = [
            ("left", ctypes.c_long),
            ("top", ctypes.c_long),
            ("right", ctypes.c_long),
            ("bottom", ctypes.c_long),
        ]

    @final
    class BITMAPINFOHEADER(ctypes.Structure):
        _fields_ = [
            ("biSize", ctypes.c_uint32),
            ("biWidth", ctypes.c_long),
            ("biHeight", ctypes.c_long),
            ("biPlanes", ctypes.c_uint16),
            ("biBitCount", ctypes.c_uint16),
            ("biCompression", ctypes.c_uint32),
            ("biSizeImage", ctypes.c_uint32),
            ("biXPelsPerMeter", ctypes.c_long),
            ("biYPelsPerMeter", ctypes.c_long),
            ("biClrUsed", ctypes.c_uint32),
            ("biClrImportant", ctypes.c_uint32),
        ]

    @final
    class BITMAPINFO(ctypes.Structure):
        _fields_ = [
            ("bmiHeader", BITMAPINFOHEADER),
            ("bmiColors", ctypes.c_uint32 * 3),
        ]

    hemf = gdi32.GetEnhMetaFileW(str(src))
    if not hemf:
        return False
    try:
        # The enhanced metafile header is >=108 bytes; use an oversized buffer
        # and read the RECTL fields by offset (bounds@8, frame@24) since the
        # exact struct layout varies by Windows version.
        hdr_buf = ctypes.create_string_buffer(256)
        if gdi32.GetEnhMetaFileHeader(hemf, 256, hdr_buf) == 0:
            return False
        import struct as _struct

        def _rect_at(off: int) -> tuple[int, int, int, int]:
            l, t, r, b = _struct.unpack_from("iiii", hdr_buf, off)
            return l, t, r, b

        # rclFrame (picture frame) is in 0.01-mm units; fall back to rclBounds.
        fl, ft, fr, fb = _rect_at(24)
        w_mm = fr - fl
        h_mm = fb - ft
        if w_mm <= 0 or h_mm <= 0:
            bl, bt, br, bb = _rect_at(8)
            w_mm = br - bl
            h_mm = bb - bt
        if w_mm <= 0 or h_mm <= 0:
            return False
        dpi = 96 * scale
        width = max(1, int(round(w_mm / 2540.0 * dpi)))
        height = max(1, int(round(h_mm / 2540.0 * dpi)))

        hdc = gdi32.CreateCompatibleDC(0)
        if not hdc:
            return False
        try:
            bmi = BITMAPINFO()
            bmi.bmiHeader.biSize = ctypes.sizeof(BITMAPINFOHEADER)
            bmi.bmiHeader.biWidth = width
            bmi.bmiHeader.biHeight = -height  # top-down
            bmi.bmiHeader.biPlanes = 1
            bmi.bmiHeader.biBitCount = 32
            bmi.bmiHeader.biCompression = 0  # BI_RGB
            ppv = ctypes.c_void_p()
            hbitmap = gdi32.CreateDIBSection(
                hdc, ctypes.byref(bmi), 0, ctypes.byref(ppv), 0, 0
            )
            if not hbitmap:
                return False
            try:
                old = gdi32.SelectObject(hdc, hbitmap)
                # White background so transparent EMF areas stay white.
                white = gdi32.GetStockObject(0)  # WHITE_BRUSH
                rect = RECTL(0, 0, width, height)
                user32.FillRect(hdc, ctypes.byref(rect), white)
                gdi32.PlayEnhMetaFile(hdc, hemf, ctypes.byref(rect))
                gdi32.SelectObject(hdc, old)

                buf = ctypes.create_string_buffer(width * height * 4)
                if (
                    gdi32.GetDIBits(hdc, hbitmap, 0, height, buf, ctypes.byref(bmi), 0)
                    == 0
                ):
                    return False
                img = PILImage.frombytes("RGBA", (width, height), bytes(buf))
                r, g, b, _a = img.split()
                img = PILImage.merge("RGB", (r, g, b))
                img.save(dst, "PNG")
                return True
            finally:
                gdi32.DeleteObject(hbitmap)
        finally:
            gdi32.DeleteDC(hdc)
    finally:
        gdi32.DeleteEnhMetaFile(hemf)


def _extract_images(doc: _Document, asset_dir: Path) -> dict[str, str]:
    """Extract inline images to *asset_dir*.

    Vector formats (EMF/WMF) are converted to PNG so they render in Markdown
    viewers; the returned ``rId`` maps to the final (PNG) filename.

    Returns ``{rId: filename}`` mapping.
    """
    image_map: dict[str, str] = {}
    if asset_dir.exists():
        for f in asset_dir.iterdir():
            if f.is_file():
                f.unlink()
    else:
        asset_dir.mkdir(parents=True, exist_ok=True)

    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        rId = rel.rId
        blob = rel.target_part.blob
        content_type = rel.target_part.content_type or "image/png"
        ext = _content_type_to_ext(content_type)
        target = asset_dir / f"{rId}{ext}"
        _ = target.write_bytes(blob)

        final_name = target.name
        if ext in (".emf", ".wmf"):
            png = asset_dir / f"{rId}.png"
            if _convert_vector_to_png(target, png):
                final_name = png.name
                target.unlink(missing_ok=True)
            else:
                print(
                    "⚠️ EMF/WMF not converted (on Windows it renders via GDI; "
                    + f"otherwise install pillow-emf / PyMuPDF): {target.name}",
                    file=sys.stderr,
                )
        image_map[rId] = final_name

    return image_map


def _content_type_to_ext(content_type: str) -> str:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
        "image/x-emf": ".emf",
        "image/x-wmf": ".wmf",
        "image/svg+xml": ".svg",
    }
    return mapping.get(content_type, ".png")


def _paragraph_image_refs(
    paragraph: Any, image_map: dict[str, str], asset_prefix: str
) -> list[str]:
    """Return ``![Image](<asset_prefix>/...)`` references for images in a paragraph.

    Detects images by the relationship-bearing leaf elements directly:
      * DrawingML ``a:blip`` with ``r:embed`` / ``r:link``
      * VML ``v:imagedata`` with ``r:id`` / ``r:embed`` (Visio / embedded OLE
        diagrams, often nested under ``w:object`` → ``v:shape`` with no
        ``w:pict`` wrapper)

    Scans the whole paragraph element via ``.iter()`` so deeply nested shapes
    are found regardless of wrapper. De-duplicated, order preserved.
    """
    refs: list[str] = []
    blip_tag = f"{_DRAW_NS}blip"
    imgd_tag = f"{_VML_NS}imagedata"
    for node in paragraph._element.iter():
        tag = node.tag
        if tag == blip_tag:
            rid = node.get(f"{_REL_NS}embed") or node.get(f"{_REL_NS}link")
            if rid and rid in image_map:
                refs.append(f"![Image]({asset_prefix}/{image_map[rid]})")
        elif tag == imgd_tag:  # VML shape (Visio / embedded diagram)
            rid = node.get(f"{_REL_NS}id") or node.get(f"{_REL_NS}embed")
            if rid and rid in image_map:
                refs.append(f"![Image]({asset_prefix}/{image_map[rid]})")

    # de-duplicate, preserve order
    seen: set[str] = set()
    out: list[str] = []
    for r in refs:
        if r not in seen:
            seen.add(r)
            out.append(r)
    return out


def _img_filename(ref: str) -> str:
    """Extract the trailing filename from a ``![Image](assets/NAME)`` ref."""
    return ref.rsplit("/", 1)[-1].rstrip(")")


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------

def _render_table(table: Any) -> list[str]:
    """Render a ``python-docx`` table as GFM Markdown table lines."""
    lines: list[str] = []
    rows = table.rows
    if not rows:
        return lines

    # Header row
    header_cells = [row.cells for row in rows]

    max_cols = max(len(row.cells) for row in rows)
    if max_cols == 0:
        return lines

    # Collect text per cell, word-wrap intelligently
    grid: list[list[str]] = []
    for row_idx, cells in enumerate(header_cells if header_cells else []):
        row_data: list[str] = []
        for cell in cells:
            cell_text = _clean_text(cell.text)
            row_data.append(cell_text)
        # Pad row to max_cols
        while len(row_data) < max_cols:
            row_data.append("")
        grid.append(row_data)

    if not grid:
        return lines

    # Compute column widths
    col_widths = [
        max(len(grid[r][c]) for r in range(len(grid)))
        for c in range(max_cols)
    ]

    # Emit header
    header_row = "| " + " | ".join(grid[0][c].ljust(col_widths[c]) for c in range(max_cols)) + " |"
    separator = "| " + " | ".join("-" * max(w, 3) for w in col_widths) + " |"
    lines.append(header_row)
    lines.append(separator)

    # Emit body
    for row_idx in range(1, len(grid)):
        body_row = "| " + " | ".join(grid[row_idx][c].ljust(col_widths[c]) for c in range(max_cols)) + " |"
        lines.append(body_row)

    return lines


# ---------------------------------------------------------------------------
# Heading level
# ---------------------------------------------------------------------------

# WordprocessingML namespace (paragraph properties / outline level)
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _heading_level(style_name: str | None, para_element: Any = None) -> int:
    """Return heading level (1-based, capped at 6) or 0 for non-headings.

    Detection order:
      1. Style name — English ``Heading N`` *and* Chinese ``标题 N``.
      2. ``w:outlineLvl`` on the paragraph — catches headings applied with a
         localized style name *or* via direct outline-level formatting.
    """
    level = 0

    # 1. Style name — English "Heading N" *and* Chinese "标题 N".
    if style_name:
        m = re.match(r"(?:heading|标题)\s*(\d+)", style_name, re.IGNORECASE)
        if m:
            return min(int(m.group(1)), 6)
        name_lower = style_name.lower().strip()
        if name_lower == "title":
            return 1
        if name_lower == "subtitle":
            return 2

    # 2. w:outlineLvl fallback — localized style names or direct formatting.
    if level == 0 and para_element is not None:
        pPr = para_element.find(f"{_W_NS}pPr")
        if pPr is None:
            return 0
        outline = pPr.find(f"{_W_NS}outlineLvl")
        if outline is None:
            return 0
        val = outline.get(f"{_W_NS}val")
        if val is None:
            return 0
        try:
            # In WordprocessingML, w:outlineLvl val 0-8 maps to heading levels
            # 1-9. val 9 means "Body Text" (not a heading) and is ignored.
            ival = int(val)
        except ValueError:
            return 0
        if 0 <= ival <= 8:
            return min(ival + 1, 6)  # outlineLvl 0 => H1

    return 0


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert(
    input_path: str | Path,
    output_path: str | Path | None = None,
    asset_dir: str | Path | None = None,
) -> tuple[Path, dict[int, list[str]]]:
    """Convert a ``.docx`` file to Markdown.

    Args:
        input_path: Path to the source ``.docx`` file.
        output_path: Path for the output ``.md`` file. If ``None``,
            defaults to ``<input_stem>.md``.
        asset_dir: Directory for extracted images. If ``None``,
            defaults to ``assets/`` next to the output file.

    Returns:
        ``(output_path, {page_number: [image_filenames, ...]})``
    """
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    doc = Document(str(input_path))

    if output_path is None:
        output_path = input_path.with_suffix(".md")
    else:
        output_path = Path(output_path)

    if asset_dir is None:
        asset_dir = output_path.parent / "assets"
    else:
        asset_dir = Path(asset_dir)

    # Relative prefix used for image refs in the Markdown output. Derived from
    # the actual (possibly custom) asset folder so refs stay correct.
    asset_prefix = asset_dir.name

    # Clear old assets
    if asset_dir.exists():
        for f in asset_dir.iterdir():
            if f.is_file():
                f.unlink()
    else:
        asset_dir.mkdir(parents=True, exist_ok=True)

    # Extract images
    image_map = _extract_images(doc, asset_dir)

    # Build page tracking for return value
    by_page: dict[int, list[str]] = {}
    current_page_images: list[str] = []

    # Output lines
    lines: list[str] = []

    # Process body elements
    # We process Document.body elements in order to handle mixed paragraphs/tables
    body = doc.element.body

    for child in body:
        tag = child.tag
        if tag.endswith("}p"):  # Paragraph
            # Find the corresponding paragraph object
            # We need to match by XML element
            para = _find_paragraph_by_element(doc.paragraphs, child)
            if para is None:
                continue

            style = para.style
            style_name = style.name if style else None
            h_level = _heading_level(style_name, child)

            # Scan for images (DrawingML + VML) on EVERY paragraph, including
            # list items and otherwise-empty paragraphs.
            img_refs = _paragraph_image_refs(para, image_map, asset_prefix)

            # Collect runs (computed once for heading / list / paragraph rendering)
            text = _runs_markdown(para.runs)

            # A paragraph styled as a heading wins over list formatting, even if
            # it also carries numbering (e.g. a "Heading 1" that is part of a
            # multilevel list). Render it as a heading first.
            if h_level > 0:
                heading_text = _runs_markdown(para.runs, keep_emphasis=False)
                marker = "#" * h_level
                lines.append(f"\n{marker} {heading_text}\n")
                for ref in img_refs:
                    lines.append(ref)
                    lines.append("")
                    if _img_filename(ref) not in current_page_images:
                        current_page_images.append(_img_filename(ref))
                continue

            # Check for list item
            list_lvl = _list_level(para)
            if list_lvl >= 0:
                is_numbered = _is_numbered_list(para)
                indent = "  " * list_lvl
                prefix = "1. " if is_numbered else "- "
                if text:
                    lines.append(f"{indent}{prefix}{text}")
                for ref in img_refs:
                    lines.append(f"{indent}{ref}")
                    if _img_filename(ref) not in current_page_images:
                        current_page_images.append(_img_filename(ref))
                lines.append("")
                continue

            # Empty paragraph
            if not text:
                if img_refs:
                    for ref in img_refs:
                        lines.append(ref)
                        if _img_filename(ref) not in current_page_images:
                            current_page_images.append(_img_filename(ref))
                    lines.append("")
                else:
                    # Empty paragraph may be intentional spacing
                    lines.append("")
                continue

            lines.append(text)
            lines.append("")
            for ref in img_refs:
                lines.append(ref)
                lines.append("")
                if _img_filename(ref) not in current_page_images:
                    current_page_images.append(_img_filename(ref))

        elif tag.endswith("}tbl"):  # Table
            table = _find_table_by_element(doc.tables, child)
            if table is None:
                continue
            table_lines = _render_table(table)
            if table_lines:
                lines.extend(table_lines)
                lines.append("")
            # Images inside table cells (emitted as a block below the table)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for ref in _paragraph_image_refs(p, image_map, asset_prefix):
                            lines.append(ref)
                            lines.append("")
                            if _img_filename(ref) not in current_page_images:
                                current_page_images.append(_img_filename(ref))

        elif tag.endswith("}sectPr"):  # Section properties (page break)
            lines.append("---")
            lines.append("")

    # Insert page anchors (every heading gets one)
    # For simplicity, we number sequentially from 1
    result_lines: list[str] = []
    para_index = 0
    for line in lines:
        if line.startswith("#"):
            para_index += 1
            result_lines.append(f"<!-- Page {para_index} -->")
            result_lines.append(line)
            if current_page_images:
                by_page[para_index] = list(current_page_images)
            current_page_images = []
        else:
            result_lines.append(line)

    # Write output
    output_path.parent.mkdir(parents=True, exist_ok=True)
    _ = output_path.write_text("\n".join(result_lines), encoding="utf-8")

    return output_path, by_page


def _find_paragraph_by_element(paragraphs: list[Any], element: Any) -> Any:
    """Find the ``python-docx`` ``Paragraph`` whose ``_element`` matches."""
    for para in paragraphs:
        if para._element is element:
            return para
    return None


def _find_table_by_element(tables: list[Any], element: Any) -> Any:
    """Find the ``python-docx`` ``Table`` whose ``_element`` matches."""
    for table in tables:
        if table._element is element:
            return table
    return None


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert Word (.docx) to Markdown with image extraction."
    )
    _ = parser.add_argument("input", type=str, help="Path to input .docx file")
    _ = parser.add_argument(
        "output", type=str, nargs="?", default=None,
        help="Path to output .md file (default: <input_stem>.md)"
    )
    _ = parser.add_argument(
        "--asset-dir", type=str, default=None,
        help="Directory for extracted images (default: <output_dir>/assets/)"
    )
    args = parser.parse_args()

    output_path, by_page = convert(args.input, args.output, args.asset_dir)
    print(f"✅ Output written to: {output_path}", file=sys.stderr)
    print(json.dumps({"md_path": str(output_path), "by_page": by_page}), file=sys.stdout)


if __name__ == "__main__":
    main()
